from docx import Document
from reportlab.pdfgen import canvas
import tempfile
import os
## REMOVED: from app.services.generator import reports

def export_report(report_id: str, format: str):
    from app.services.generator import preview
    print(f"🔄 Exporting report {report_id} in {format} format")
    
    # Use the preview function to get the existing report
    result = preview(report_id)
    if "error" in result:
        print(f"❌ Export failed: {result['error']}")
        return {"error": f"Report not found: {result['error']}"}
    
    report = result["preview"]
    print(f"✅ Report data retrieved successfully")
    # Format the sections for export
    content = {}
    for section in report["sections"]:
        title = section["title"]
        text_content = ""
        for item in section["content"]:
            # Clean text content to handle Unicode properly
            clean_text = str(item["text"]).encode('utf-8', errors='ignore').decode('utf-8')
            text_content += clean_text + "\n\n"
            if item.get("citations"):
                text_content += "Citations:\n"
                for citation in item["citations"]:
                    clean_snippet = str(citation.get('snippet', '')).encode('utf-8', errors='ignore').decode('utf-8')
                    text_content += f"- {citation.get('source_id')}, Page {citation.get('page')}: {clean_snippet[:100]}...\n"
        content[title] = text_content
    if format == "docx":
        doc = Document()
        
        # Add title
        title = doc.add_heading('AI Generated Report', 0)
        
        # Add report ID
        doc.add_paragraph(f'Report ID: {report_id}')
        doc.add_paragraph('')  # Empty line
        
        for section, text in content.items():
            # Add section heading
            doc.add_heading(section, level=1)
            
            # Add section content with proper formatting
            lines = text.splitlines()
            for line in lines:
                if line.strip():  # Skip empty lines
                    # Handle citations specially
                    if line.startswith("- ") and "Page" in line:
                        # This is a citation line
                        p = doc.add_paragraph(line, style='List Bullet')
                    else:
                        # Regular content
                        doc.add_paragraph(line)
                else:
                    # Add empty paragraph for spacing
                    doc.add_paragraph('')
        
        tmp_path = tempfile.mktemp(suffix=".docx")
        doc.save(tmp_path)
        with open(tmp_path, "rb") as f:
            data = f.read()
        os.remove(tmp_path)
        print(f"✅ DOCX export completed, file size: {len(data)} bytes")
        return {"report_id": report_id, "format": format, "file": data}
    elif format == "pdf":
        tmp_path = tempfile.mktemp(suffix=".pdf")
        c = canvas.Canvas(tmp_path, pagesize=(612, 792))  # Letter size
        width, height = 612, 792
        margin = 50
        y = height - margin
        
        for section, text in content.items():
            # Add section title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(margin, y, section)
            y -= 30
            
            # Add section content with proper text wrapping
            c.setFont("Helvetica", 12)
            lines = text.splitlines()
            for line in lines:
                # Handle long lines by wrapping them
                if len(line) > 80:  # Approximate characters per line
                    words = line.split()
                    current_line = ""
                    for word in words:
                        if len(current_line + " " + word) > 80:
                            if current_line:
                                c.drawString(margin + 20, y, current_line)
                                y -= 15
                            current_line = word
                        else:
                            current_line += " " + word if current_line else word
                    if current_line:
                        c.drawString(margin + 20, y, current_line)
                        y -= 15
                else:
                    c.drawString(margin + 20, y, line)
                    y -= 15
                
                # Check if we need a new page
                if y < margin:
                    c.showPage()
                    y = height - margin
            
            y -= 20  # Space between sections
            
            # Check if we need a new page before next section
            if y < margin + 50:
                c.showPage()
                y = height - margin
        
        c.save()
        with open(tmp_path, "rb") as f:
            data = f.read()
        os.remove(tmp_path)
        print(f"✅ PDF export completed, file size: {len(data)} bytes")
        return {"report_id": report_id, "format": format, "file": data}
    else:
        return {"error": "Unsupported export format"}
